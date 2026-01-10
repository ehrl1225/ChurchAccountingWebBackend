import docx
from docx.document import Document
from docx.styles.style import BaseStyle, ParagraphStyle
from docx.types import annotations
from docx.styles.styles import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.table import _Cell, Table
from datetime import datetime
import calendar
from typing import Optional

def setCellBorder(cell: _Cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def setCellBorders(cells: list[_Cell], **kwargs):
    for cell in cells:
        setCellBorder(cell, **kwargs)

def getRowCells(table: Table, row:int, columns: list[int]) -> list[_Cell]:
    cells = []
    for column in columns:
        cell:_Cell = table.cell(row, column)
        cells.append(cell)
    return cells

def setCellColor(cell: _Cell, color: str):
    cell_xml_element = cell._tc
    table_cell_properties = cell_xml_element.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    shade_obj.set(qn('w:fill'), color)
    table_cell_properties.append(shade_obj)

def mergeCell(cell1: _Cell, cell2: _Cell):
    cell1.merge(cell2)

def getMonthRange(year, month) -> tuple[int, int]:
    return calendar.monthrange(year, month)

def amountToString(amount:int) -> str:
    data = ""
    str_amount = str(amount)
    str_amount_len = len(str_amount)-1
    for i in range(str_amount_len+1 if amount >= 0 else str_amount_len):

        if i%3 == 0 and i != 0:
            data = ',' + data
        data = str_amount[str_amount_len-i] + data
    return data if amount >= 0 else '-'+data

def setFont(doc:Document, name:str, font_name:str, font_size:int, bold:bool) -> WD_STYLE_TYPE:
    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.font.size = Pt(font_size)
    style.font.bold = bold
    return style

def setCellText(cell: _Cell, text:str, style:Optional[WD_STYLE_TYPE], align:Optional[WD_ALIGN_PARAGRAPH], vertical_align:Optional[WD_ALIGN_VERTICAL]):
    cell.paragraphs[0].text = text
    if style is not None:
        cell.paragraphs[0].style = style
    if align is not None:
        cell.paragraphs[0].alignment = align
    if vertical_align is not None:
        cell.vertical_alignment = vertical_align

single = {"val":"single"}
single_bold = {"val": "single", "sz": 10}
double = {"val": "double"}
cell_color = "#E5E5E5"
last_cell_color = "#BFBFBF"