from docx.shared import Cm

from domain.organization.organization.repository import OrganizationRepository
from .word_util import *
from docx.enum.table import WD_TABLE_ALIGNMENT
import io

from sqlalchemy.ext.asyncio import AsyncSession

from common.enum.tx_type import TxType
from domain.file.word.dto.create_settlement_dto import CreateSettlementDto
from domain.ledger.receipt.dto import ReceiptSummaryParams, ReceiptSummaryDto, ReceiptSummaryCategoryDto, \
    ReceiptSummaryItemDto
from domain.ledger.receipt.service import ReceiptService


class WordService:

    def __init__(
            self,
            organization_repository: OrganizationRepository,
            receipt_service: ReceiptService
    ):
        self.organization_repository = organization_repository
        self.receipt_service = receipt_service

    async def create_month_document(self,db:AsyncSession, create_settlement:CreateSettlementDto):
        doc:Document = docx.Document()

        section = doc.sections[0]
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

        # 기본 설정

        style_1 = setFont(doc, "style_1", "맑은 고딕", 15, True)
        style_2 = setFont(doc, "style_2", "맑은 고딕", 10, False)
        style_3 = setFont(doc, "style_3", "맑은 고딕", 13, True)
        style_4 = setFont(doc, "style_4", "맑은 고딕", 30, False)
        style_5 = setFont(doc, "style_5", "맑은 고딕", 15, False)

        # 데이터

        year = create_settlement.year
        month = create_settlement.month_number
        first_day, last_day = calendar.monthrange(year, month)

        organization = await self.organization_repository.find_by_id(db, create_settlement.organization_id)
        organization_name = organization.name

        data:ReceiptSummaryDto = await self.receipt_service.get_summary_receipt(db, receipt_summary_params=ReceiptSummaryParams(
            summary_type=create_settlement.summary_type,
            month_number=create_settlement.month_number,
            organization_id=create_settlement.organization_id,
            year=create_settlement.year,
            event_id=create_settlement.event_id,
        ))
        income_categories = []
        outcome_categories = []
        income_items_count = 0
        outcome_items_count = 0
        income_total = data.total_income
        outcome_total = data.total_outcome
        balance = data.balance

        for category in data.categories:
            if category.tx_type == TxType.INCOME:
                income_categories.append(category)
                income_items_count += len(category.items)
            else:
                outcome_categories.append(category)
                outcome_items_count += len(category.items)

        total_row_count = income_items_count + outcome_items_count + 6
        income_start_row_index = 1
        income_total_row_index = income_items_count + 1
        outcome_start_row_index = income_items_count + 2
        outcome_total_row_index = income_items_count + outcome_items_count + 2

        total_sum_row_index = outcome_total_row_index + 2

        # 작성 시작

        p1 = doc.add_paragraph(f'{year}년도 {organization_name} {month}월 결산안', style=style_1)
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p2 = doc.add_paragraph(f"({year}.{month}.{first_day} ~ {year}.{month}.{last_day})", style=style_2)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p3 = doc.add_paragraph("[단위 : 원]", style=style_2)
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        table = doc.add_table(rows=total_row_count, cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        first_line_items = ["구 분", "항 목", "내 용", "금 액", "비 고"]
        mergeCell(table.cell(0, 3), table.cell(0, 4))
        for col in [i for i in range(6) if i != 4]:
            cell = table.cell(0, col)
            if col == 0:
                setCellBorder(cell, start=single_bold)
            elif col == 5:
                setCellBorder(cell, start=single, end=single_bold)
            else:
                setCellBorder(cell, start=single)
            setCellBorder(cell, top=single_bold, bottom=single_bold)
            setCellColor(cell, cell_color)
            if col < 4:
                setCellText(cell, first_line_items[col], style_3, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_VERTICAL.CENTER)
            else:
                setCellText(cell, first_line_items[4], style_3, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_VERTICAL.CENTER)
        self.setCategories(
            table,
            style_3,
            "수\n\n입",
            income_categories,
            income_start_row_index,
            income_items_count,
            income_total_row_index,
            income_total
        )

        self.setCategories(
            table,
            style_3,
            "지\n\n출",
            outcome_categories,
            outcome_start_row_index,
            outcome_items_count,
            outcome_total_row_index,
            outcome_total
        )

        income_total_sum_cell = table.cell(total_sum_row_index, 0)
        income_total_sum_be_merge_cell = table.cell(total_sum_row_index, 1)
        income_total_sum_amount_cell = table.cell(total_sum_row_index + 1, 0)
        income_total_sum_amount_be_merge_cell = table.cell(total_sum_row_index + 1, 1)

        mergeCell(income_total_sum_cell, income_total_sum_be_merge_cell)
        mergeCell(income_total_sum_amount_cell, income_total_sum_amount_be_merge_cell)

        setCellBorder(income_total_sum_cell, start=single_bold, end=single, top=single_bold, bottom=single)
        setCellColor(income_total_sum_cell, last_cell_color)
        setCellBorder(income_total_sum_amount_cell, start=single_bold, end=single, bottom=single_bold)

        setCellText(income_total_sum_cell, f"{month}월 수입", style_3, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_VERTICAL.CENTER)

        setCellText(income_total_sum_amount_cell, amountToString(income_total), style_3, WD_ALIGN_PARAGRAPH.CENTER,
                    WD_ALIGN_VERTICAL.CENTER)

        outcome_total_sum_cell = table.cell(total_sum_row_index, 2)
        outcome_total_sum_be_merge_cell = table.cell(total_sum_row_index, 3)

        outcome_total_sum_amount_cell = table.cell(total_sum_row_index + 1, 2)
        outcome_total_sum_amount_be_merge_cell = table.cell(total_sum_row_index + 1, 3)

        setCellBorder(outcome_total_sum_cell, start=single, end=single, top=single_bold, bottom=single)
        setCellColor(outcome_total_sum_cell, last_cell_color)
        setCellBorder(outcome_total_sum_amount_cell, start=single, end=single, bottom=single_bold)

        setCellText(outcome_total_sum_cell, f"{month}월 지출", style_3,
                    WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_VERTICAL.CENTER)

        setCellText(outcome_total_sum_amount_cell,
                    amountToString(outcome_total),
                    style_3,
                    WD_ALIGN_PARAGRAPH.CENTER,
                    None)

        mergeCell(outcome_total_sum_cell, outcome_total_sum_be_merge_cell)
        mergeCell(outcome_total_sum_amount_cell, outcome_total_sum_amount_be_merge_cell)

        remain_total_sum_cell = table.cell(total_sum_row_index, 4)
        remain_total_sum_be_merge_cell = table.cell(total_sum_row_index, 5)

        remain_total_sum_amount_cell = table.cell(total_sum_row_index + 1, 4)
        remain_total_sum_amount_be_merge_cell = table.cell(total_sum_row_index + 1, 5)

        setCellBorder(remain_total_sum_cell, start=single, end=single_bold, top=single_bold, bottom=single)
        setCellColor(remain_total_sum_cell, last_cell_color)
        setCellBorder(remain_total_sum_amount_cell, start=single, end=single, bottom=single_bold)

        setCellText(remain_total_sum_cell,
                    f"{month}월 잔액",
                    style_3,
                    WD_ALIGN_PARAGRAPH.CENTER,
                    WD_ALIGN_VERTICAL.CENTER)

        setCellText(remain_total_sum_amount_cell,
                    amountToString(balance),
                    style_3,
                    WD_ALIGN_PARAGRAPH.CENTER,
                    WD_ALIGN_VERTICAL.CENTER)

        mergeCell(remain_total_sum_cell, remain_total_sum_be_merge_cell)
        mergeCell(remain_total_sum_amount_cell, remain_total_sum_amount_be_merge_cell)

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.0

        doc.add_paragraph(" ")
        doc.paragraphs[3].runs[0].add_break(docx.text.run.WD_BREAK.PAGE)

        table2 = doc.add_table(rows=1, cols=5)
        table2.alignment = WD_TABLE_ALIGNMENT.CENTER

        confirm_date_cell = table2.cell(0, 0)
        confirm_content_cell = table2.cell(0, 1)
        confirm_church_cell = table2.cell(0, 4)
        for i in range(2, 4):
            confirm_content_be_merge_cell = table2.cell(0, i)
            mergeCell(confirm_content_cell, confirm_content_be_merge_cell)

        setCellBorder(confirm_date_cell, start=single, end=single, top=single, bottom=single)
        setCellBorder(confirm_content_cell, end=single, top=single, bottom=single)
        setCellBorder(confirm_church_cell, end=single, top=single, bottom=single)

        setCellText(confirm_date_cell, f"{year}-{month:02d}월", style_5, None, WD_ALIGN_VERTICAL.CENTER)

        setCellText(confirm_content_cell, "증빙내용", style_4, WD_ALIGN_PARAGRAPH.CENTER, None)

        setCellText(confirm_church_cell, f"{organization_name}", None, WD_ALIGN_PARAGRAPH.CENTER,
                    WD_ALIGN_VERTICAL.CENTER)

        doc.add_paragraph(" ")

        table3 = doc.add_table(rows=8, cols=6)

        put_receipt_here_cell = table3.cell(0, 0)
        setCellText(put_receipt_here_cell, "영수증 붙이는 곳", None, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_VERTICAL.CENTER)

        for i in range(8):
            receipt_cell = table3.cell(i, 0)
            for j in range(1, 3):
                receipt_be_merge_cell = table3.cell(i, j)
                mergeCell(receipt_cell, receipt_be_merge_cell)

        pay_list_cell = table3.cell(0, 3)

        for i in range(4, 6):
            pay_list_be_merge_cell = table3.cell(0, i)
            mergeCell(pay_list_cell, pay_list_be_merge_cell)

        setCellText(pay_list_cell, "증빙 없는 지출 내역", None, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_VERTICAL.CENTER)

        content_cell = table3.cell(1, 3)
        setCellText(content_cell, "내용", None, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_VERTICAL.CENTER)
        setCellBorder(content_cell, start=single, end=single, top=single, bottom=single)
        setCellColor(content_cell, cell_color)

        pay_to_cell = table3.cell(1, 4)
        setCellText(pay_to_cell, "지출처", None, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_VERTICAL.CENTER)
        setCellBorder(pay_to_cell, end=single, top=single, bottom=single)
        setCellColor(pay_to_cell, cell_color)

        amount_cell = table3.cell(1, 5)
        setCellText(amount_cell, "금액", None, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_VERTICAL.CENTER)
        setCellBorder(amount_cell, end=single, top=single, bottom=single)
        setCellColor(amount_cell, cell_color)

        for i in range(2, 8):
            current_content_cell = table3.cell(i, 3)
            setCellBorder(current_content_cell, start=single, end=single, top=single, bottom=single)
            for j in range(4, 6):
                current_content_side_cell = table3.cell(i, j)
                setCellBorder(current_content_side_cell, end=single, top=single, bottom=single)

        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        return stream

    @staticmethod
    def setCategories(
            table: Table,
            style:WD_STYLE_TYPE,
            start_text: str,
            categories: list[ReceiptSummaryCategoryDto],
            start_row_index: int,
            items_count: int,
            total_row_index: int,
            total:int
    ):
        cell = table.cell(start_row_index, 0)
        setCellText(cell, start_text, style, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_VERTICAL.CENTER)
        if items_count >= 2:
            setCellBorder(cell, start=single_bold, end=single)
            setCellColor(cell, cell_color)

            for index in range(1, items_count):
                row_index = start_row_index + index
                be_merge_cell = table.cell(row_index, 0)

                if index == items_count - 1:
                    setCellBorder(be_merge_cell, start=single_bold, end=single, bottom=double)
                else:
                    setCellBorder(be_merge_cell, start=single_bold, end=single)
                mergeCell(cell, be_merge_cell)
        else:
            setCellBorder(cell, start=single_bold, end=single, bottom=double)
            setCellColor(cell, cell_color)

        category_row_index = start_row_index
        category_count = len(categories)
        for category_index, category in enumerate(categories):
            category_cell = table.cell(category_row_index, 1)
            category_total_amount_cell = table.cell(category_row_index, 4)
            items:list[ReceiptSummaryItemDto] = category.items
            item_count = len(items)

            for item_index in range(item_count):
                item_row_index = category_row_index + item_index

                current_item_etc_cell = table.cell(item_row_index, 5)
                if item_index == item_count - 1:
                    if category_index == category_count - 1:
                        setCellBorders(getRowCells(table, item_row_index, [1, 2, 3, 4]), end=single, bottom=double)
                        setCellBorder(current_item_etc_cell, end=single_bold, bottom=double)
                    else:
                        setCellBorders(getRowCells(table, item_row_index, [1, 2, 3, 4]), end=single, bottom=single)
                        setCellBorder(current_item_etc_cell, end=single_bold, bottom=single)
                else:
                    setCellBorders(getRowCells(table, item_row_index, [1, 4]), end=single)
                    setCellBorders(getRowCells(table, item_row_index, [2, 3]), end=single, bottom=single)
                    setCellBorder(current_item_etc_cell, end=single_bold, bottom=single)

                category_be_merge_cell = table.cell(item_row_index, 1)
                if item_index != 0:
                    mergeCell(category_cell, category_be_merge_cell)

                category_total_amount_be_merge_cell = table.cell(item_row_index, 4)
                if item_index != 0:
                    mergeCell(category_total_amount_cell, category_total_amount_be_merge_cell)

            setCellText(category_cell, category.category_name, None, WD_ALIGN_PARAGRAPH.CENTER,
                        WD_ALIGN_VERTICAL.CENTER)

            for item_index, item in enumerate(category.items):
                item_row_index = category_row_index + item_index
                item_cell = table.cell(item_row_index, 2)
                setCellText(item_cell, item.item_name, None, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_VERTICAL.CENTER)

                item_amount_cell = table.cell(item_row_index, 3)
                setCellText(item_amount_cell, amountToString(item.amount), None,
                            WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_VERTICAL.CENTER)

                item_etc_cell = table.cell(item_row_index, 5)
                setCellText(item_etc_cell, "", None, WD_ALIGN_PARAGRAPH.CENTER,
                            WD_ALIGN_VERTICAL.CENTER)

            category_total_amount_cell = table.cell(category_row_index, 4)
            setCellText(category_total_amount_cell, amountToString(category.amount), None,
                        WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_VERTICAL.CENTER)

            category_row_index += item_count

        total_cell = table.cell(total_row_index, 0)
        setCellBorder(total_cell, start=single_bold, end=single, bottom=single_bold)
        setCellColor(total_cell, cell_color)
        for total_be_merge_index in range(1, 3):
            total_be_merge_cell = table.cell(total_row_index, total_be_merge_index)
            mergeCell(total_cell, total_be_merge_cell)

        setCellText(total_cell, "합 계", style, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_VERTICAL.CENTER)
        total_amount_cell = table.cell(total_row_index, 3)
        total_amount_be_merge_cell = table.cell(total_row_index, 4)

        setCellBorder(total_amount_cell, end=single, bottom=single_bold)
        setCellColor(total_amount_cell, cell_color)

        mergeCell(total_amount_cell, total_amount_be_merge_cell)
        setCellText(total_amount_cell, amountToString(total), style, WD_ALIGN_PARAGRAPH.CENTER,
                    WD_ALIGN_VERTICAL.CENTER)

        total_etc_cell = table.cell(total_row_index, 5)
        setCellBorder(total_etc_cell, end=single_bold, bottom=single_bold)
        setCellColor(total_etc_cell, cell_color)
